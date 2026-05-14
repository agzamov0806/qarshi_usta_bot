from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
SRC_MD = ROOT / "docs" / "qarshi-usta-bot-tz.md"
OUT_DOCX = ROOT / "docs" / "qarshi-usta-bot-tz.docx"
OUT_PDF = ROOT / "docs" / "qarshi-usta-bot-tz.pdf"


def _read_md_lines() -> list[str]:
    return SRC_MD.read_text(encoding="utf-8").splitlines()


def export_docx(lines: list[str]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for ln in lines:
        s = ln.rstrip()
        if not s:
            doc.add_paragraph("")
            continue

        if s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
            continue

        if s.startswith("- "):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
            continue

        if s.startswith("  - "):
            p = doc.add_paragraph(s[4:].strip(), style="List Bullet 2")
            _ = p
            continue

        doc.add_paragraph(s)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def export_pdf(lines: list[str]) -> None:
    OUT_PDF.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11.5, leading=14)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14)

    story: list = []
    for ln in lines:
        s = ln.rstrip()
        if not s:
            story.append(Spacer(1, 6))
            continue

        if s.startswith("# "):
            story.append(Paragraph(s[2:].strip(), h1))
            story.append(Spacer(1, 6))
            continue
        if s.startswith("## "):
            story.append(Paragraph(s[3:].strip(), h2))
            story.append(Spacer(1, 4))
            continue
        if s.startswith("### "):
            story.append(Paragraph(s[4:].strip(), h3))
            story.append(Spacer(1, 3))
            continue

        if s.startswith("- "):
            story.append(Paragraph(f"• {s[2:].strip()}", body))
            continue
        if s.startswith("  - "):
            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;• {s[4:].strip()}", body))
            continue

        story.append(Paragraph(s.replace("`", ""), body))

    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Qarshi Usta Bot — Texnik topshiriq",
    )
    doc.build(story)


def main() -> None:
    if not SRC_MD.exists():
        raise SystemExit(f"Source markdown not found: {SRC_MD}")
    lines = _read_md_lines()
    export_docx(lines)
    export_pdf(lines)
    print(f"OK: {OUT_DOCX}")
    print(f"OK: {OUT_PDF}")


if __name__ == "__main__":
    main()

